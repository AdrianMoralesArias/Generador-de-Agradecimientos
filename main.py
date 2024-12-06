from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

os.system("cls")

def crear_hoja_agradecimiento(nombre):
    # Crear un documento de Word
    doc = Document()

    # Agregar el título
    titulo = doc.add_heading(level=1)
    run_titulo = titulo.add_run("Carta de Agradecimiento")
    run_titulo.font.size = Pt(20)
    run_titulo.font.bold = True
    run_titulo.font.color.rgb = RGBColor(0, 102, 204)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Agregar un espacio
    doc.add_paragraph("\n")

    # Cuerpo del mensaje
    parrafo = doc.add_paragraph()
    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alineación a la izquierda
    run_parrafo = parrafo.add_run(
        f"Por este medio, extendemos nuestro más profundo y sincero agradecimiento a {nombre} por "
        f"su valiosa participación en la 'Actividad de Ejemplo'. Su contribución y apoyo han sido fundamentales "
        f"para el éxito de este evento, el cual tiene como objetivo fomentar la colaboración y el crecimiento "
        f"mutuo entre todos los participantes. Gracias a su dedicación y esfuerzo, pudimos llevar a cabo esta "
        f"actividad con un alto nivel de calidad, lo que dejó un impacto positivo en todos los asistentes.\n\n"
        
        f"Su disposición para participar y compartir en este importante evento refleja un compromiso encomiable "
        f"con los valores que promovemos, tales como la solidaridad, el trabajo en equipo y el desarrollo continuo. "
        f"Su entusiasmo y participación activa no solo enriquecieron la experiencia, sino que también demostraron "
        f"su profunda convicción en la importancia de contribuir al bienestar colectivo. Apreciamos enormemente el "
        f"tiempo y los recursos que dedicó, y su apoyo fue clave para que el evento alcanzara los resultados esperados.\n\n"
        
        f"Gracias a su valiosa participación, logramos fortalecer la cooperación entre todos los participantes, "
        f"fomentando un ambiente de aprendizaje y crecimiento mutuo. Los objetivos del evento, que incluían la "
        f"creación de redes de colaboración y el intercambio de conocimientos, fueron alcanzados gracias a la "
        f"dedicación de personas como usted. Además, su participación dejó una huella positiva que perdurará en "
        f"nuestro esfuerzo por seguir creando oportunidades de desarrollo.\n\n"
    )
    run_parrafo.font.size = Pt(12)

    # Agregar un espacio
    doc.add_paragraph("\n")

    # Firma
    firma = doc.add_paragraph()
    firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_firma = firma.add_run(
        "\n\n\n______________________________\nAdrian Morales Arias\nEncargado del Evento"
    )
    run_firma.font.size = Pt(12)

    # Guardar el documento
    filename = f"Agradecimiento_{nombre.replace(' ', '_')}.docx"
    doc.save(filename)
    print(f"Hoja de agradecimiento creada: {filename}")

if __name__ == "__main__":
    print("Generador de Hojas de Agradecimiento")
    print("Ingrese los nombres de las personas o empresas separados por comas. Escriba 'FIN' para terminar.")

    # Solicitar la entrada de nombres
    nombres_input = input("Nombres de las personas o empresas: ").strip()

    if nombres_input.upper() == "FIN":
        print("Saliendo del programa.")
    else:
        # Separar los nombres por comas
        nombres = [nombre.strip() for nombre in nombres_input.split(',')]

        if not nombres:
            print("No se ingresaron nombres. Saliendo del programa.")
        else:
            print("Generando hojas de agradecimiento...")
            for nombre in nombres:
                crear_hoja_agradecimiento(nombre)
            print("Todas las hojas de agradecimiento han sido creadas.")
